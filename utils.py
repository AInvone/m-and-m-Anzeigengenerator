import os
import re
from datetime import datetime
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from openai import OpenAI

# ---------- LLM Functions ----------

def generate_real_estate_ad_deeps(args, prompt):
    client = OpenAI(
        api_key=args['deepseek_api_key'],
        base_url=args['deepseek_base_url'],
    )
    response = client.chat.completions.create(
        model=args['deepseek_model'],
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content


def create_prompt(params):
    return f"""
Erstelle ein professionelles Immobilien-Exposé für Immoscout24 auf Deutsch mit folgender Struktur und Formatierung:

1. **Titel**: Eine aufmerksamkeitsstarke Überschrift.
2. **Untertitel**: Ein kurzer, inspirierender Untertitel.
3. **Objektübersicht**: Kurzbeschreibung der Immobilie (2-3 Sätze).
4. **Key Facts**: Tabelle mit den wichtigsten Eckdaten (Kaufpreis, Wohnfläche, Grundstück, Zimmer, Baujahr, Haustyp, Energieklasse, etc.).
5. **Ausführliche Beschreibung**: Mindestens 4 Abschnitte, jeweils mit Überschrift und ausführlichem Fließtext. Gehe auf Lage, Ausstattung, Besonderheiten, Zielgruppe, Vorteile und Umgebung ein.
6. **Ausstattung & Merkmale**: Bullet-Liste der wichtigsten Ausstattungsmerkmale.
7. **Lagebeschreibung**: Ausführlicher Abschnitt zur Lage.
8. **Energieinformationen**: Tabelle mit Energieausweis-Daten (Energieklasse, Verbrauch, Baujahr, Heizungsart, etc.).
9. **Kontakt**: Platzhalter für Kontaktdaten.
10. **Bildergalerie**: Platzhalter für Bilder.

Format:
- Nutze Markdown für Überschriften (##, ###), Tabellen und Listen.
- Jede Sektion klar abtrennen.
- Keine Kommentare oder Anmerkungen außerhalb des Exposé-Textes.

Daten:
Aktion: {params['action']}
Immobilientyp: {params['immotype']}
Ort: {params['location']}
Adresse: {params.get('address', '')}
Preis: {params['price']}
Wohnfläche: {params['living_area']}
Grundstücksgröße: {params['lot_size']}
Zimmeranzahl: {params['rooms']}
Baujahr: {params['year_built']}
Haustyp: {params['house_type']}
Heizungsart: {params['heating']}
Internetgeschwindigkeit: {params['internet_speed']}
Energieeffizienzklasse: {params['energy_efficiency_class']}
Ausstattung/Merkmale: {params['features']}
Energieverbrauch: {params.get('energy_consumption', '')}
Kontakt: {params.get('contact', '')}

"""

def trim_trailing_notes(text: str) -> str:
    lines = text.strip().split("\n")
    while lines and (
        "diese anzeige" in lines[-1].lower()
        or lines[-1].strip() in ("---", "")
    ):
        lines.pop()
    return "\n".join(lines)

# ---------- Markdown Cleaning ----------

def _process_markdown_line(line):
    parts = []
    buffer = []
    in_bold = False
    for char in line:
        if char == "*" and len(buffer) >= 1 and buffer[-1] == "*":
            buffer.pop()
            if in_bold:
                parts.append(('bold', ''.join(buffer)))
                in_bold = False
            else:
                parts.append(('text', ''.join(buffer)))
                in_bold = True
            buffer = []
        else:
            buffer.append(char)
    if buffer:
        parts.append(('text', ''.join(buffer)))
    return parts

def clean_heading(text):
    # Remove leading #, whitespace, and possible Markdown heading artifacts
    return re.sub(r'^#+\s*', '', text).strip()

def clean_markdown(text):
    lines = []
    first_line = True
    for line in text.split("\n"):
        line = line.strip()
        # Remove code block markers and LLM placeholders
        if not line or line.startswith("```") or "[Platzhalter für Bilder]" in line or "Fügen Sie hier hochwertige Fotos" in line or line == "---":
            continue
        # HEADING 1: erste fette Zeile als Haupttitel
        if first_line and line.startswith("**") and line.endswith("**"):
            clean_title = line.strip("*").strip()
            lines.append(("heading1", [("text", clean_title)]))
            first_line = False
            continue
        first_line = False
        # Markdown-Tabelle
        if "|" in line and re.match(r"^\|.*\|$", line):
            raw_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(set(c) <= {"-"} for c in raw_cells):
                continue
            lines.append(("table_row", raw_cells))
            continue
        # Unterüberschriften
        if line.startswith("###"):
            lines.append(("heading2", clean_heading(line.replace("###", ""))))
            continue
        elif line.startswith("**") and line.endswith("**"):
            line_clean = re.sub(r"\*+", "", line).strip()
            lines.append(("heading3", clean_heading(line_clean)))
            continue
        # Bullet (aber keine kaputten wie • --)
        bullet_match = re.match(r'^\s*[-•✔🔹]\s+(.+)', line)
        if bullet_match and bullet_match.group(1).strip():
            lines.append(("bullet", (1, _process_markdown_line(bullet_match.group(1)))))
            continue
        # Fließtext
        processed = _process_markdown_line(line)
        if any(t[0] == 'bold' for t in processed):
            lines.append(("bold_text", processed))
        else:
            lines.append(("paragraph", processed))
    return lines


def set_table_border(tbl, border_dir, val="single", size="4", color="888888"):
    """
    Setzt Tabellenrahmen in einem docx-Tabellelement.
    """
    tbl_pr = tbl.tblPr

    # Try to find existing tblBorders element 
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)


    # Create and configure the border element 
    border = OxmlElement(f"w:{border_dir}")
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)

    # Add or replace the specific border direction 
    existing_border = tbl_borders.find(qn(f"w:{border_dir}"))
    if existing_border is not None:
        tbl_borders.remove(existing_border)
    tbl_pr.append(border)


def add_clean_table_to_docx(doc, rows, bold_header=True):
    """
    Fügt eine modern formatierte Tabelle mit grauen Rändern und sauberen Zellen in ein docx-Dokument ein.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    if not rows or not all(isinstance(r, list) for r in rows):
        return

    # Filtere alle Zeilen raus, die nur aus '---' bestehen
    filtered_rows = [
        r for r in rows
        if not all(set(cell.strip()) <= {"-"} for cell in r)
    ]
    if not filtered_rows:
        return

    table = doc.add_table(rows=0, cols=len(filtered_rows[0]))
    table.style = "Table Grid"

    for direction in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        set_table_border(table._tbl, direction)

    for row_idx, row_cells in enumerate(filtered_rows):
        row = table.add_row()
        for col_idx, cell_text in enumerate(row_cells):
            clean_text = cell_text.replace("**", "").strip()
            is_bold = (bold_header and row_idx == 0)

            cell = row.cells[col_idx]
            para = cell.paragraphs[0]
            run = para.add_run(clean_text)
            run.bold = is_bold
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.space_after = Inches(0.05)

    return table


# ---------- Image Helper ----------

def add_logo_top_right(doc, logo_path):
    if logo_path and os.path.exists(logo_path):
        section = doc.sections[0]
        section.top_margin = Inches(1.2)
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.add_run().add_picture(logo_path, width=Inches(2))

def rescale_img(input_path_or_folder, output_folder="data/rescaled_details", max_width=1500, quality=100):
    """
    Rescales a single image or all images in a folder to max width.
    Saves them as optimized JPEGs in `output_folder` and returns paths.

    Args:
        input_path_or_folder (str): Path to image file or directory.
        output_folder (str): Where to save rescaled images.
        max_width (int): Max pixel width.
        quality (int): JPEG quality (default 85).

    Returns:
        str | dict: Rescaled image path (if single image) or dict of {original_path: resized_path}
    """
    os.makedirs(output_folder, exist_ok=True)

    def _rescale_and_save(img_path):
        try:
            with Image.open(img_path) as img:
                img_format = img.format or "JPEG"
                width, height = img.size

                if width > max_width:
                    new_height = int(max_width * height / width)
                    img = img.resize((max_width, new_height), Image.LANCZOS)

                base_name = os.path.splitext(os.path.basename(img_path))[0] + ".jpg"
                out_path = os.path.join(output_folder, base_name)
                img.convert("RGB").save(out_path, format="JPEG", quality=quality, optimize=True)
                return out_path
        except Exception as e:
            print(f"⚠️ Fehler beim Skalieren von {img_path}: {e}")
            return img_path  # fallback

    if os.path.isfile(input_path_or_folder):
        return _rescale_and_save(input_path_or_folder)
    
    elif os.path.isdir(input_path_or_folder):
        results = {}
        for f in os.listdir(input_path_or_folder):
            full_path = os.path.join(input_path_or_folder, f)
            if f.lower().endswith((".png", ".jpg", ".jpeg")):
                resized = _rescale_and_save(full_path)
                results[full_path] = resized
        return results

    else:
        raise ValueError(f"Pfad nicht gefunden: {input_path_or_folder}")
    
def add_main_image(doc, path, caption=""):
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run().add_picture(path, width=Inches(5.5))
        if caption:
            caption_paragraph = doc.add_paragraph(caption, style="Caption")
            caption_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def clean_filename_for_caption(fname):
    name = os.path.splitext(os.path.basename(fname))[0]
    return name.replace("-", " ").replace("_", " ").title()

def classify_room_priority(name):
    name = name.lower()
    MAIN = ["wohnzimmer", "flur", "schlafzimmer", "kinderzimmer"]
    FUNC = ["küche", "bad", "balkon", "garage", "wc"]
    for i, r in enumerate(MAIN):
        if r in name:
            return (0, i)
    for i, r in enumerate(FUNC):
        if r in name:
            return (1, i)
    return (2, 99)

def extract_number(fname):
    match = re.search(r'(\d+)(?=\.\w+$)', fname)
    return int(match.group(1)) if match else 999

def add_image_gallery_from_folder(doc, folder, title="Innenansichten", per_row=2, width=2.5):
    if not os.path.isdir(folder):
        return
    images = [os.path.join(folder, f) for f in os.listdir(folder)
              if f.lower().endswith((".png", ".jpg", ".jpeg"))]
    images.sort(key=lambda p: (classify_room_priority(p), extract_number(p)))
    if not images:
        return

    doc.add_heading(title, level=2)
    for i in range(0, len(images), per_row):
        table = doc.add_table(rows=1, cols=per_row)
        row = table.rows[0].cells
        for j in range(per_row):
            if i + j < len(images):
                img = images[i + j]
                p = row[j].paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run().add_picture(img, width=Inches(width))
                row[j].add_paragraph(clean_filename_for_caption(img)).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# ---------- Save Word File ----------

def load_agb_text():
    agb_path = os.path.join(os.path.dirname(__file__), "agb.txt")
    with open(agb_path, "r", encoding="utf-8") as f:
        return f.read()

def add_call_to_action(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F7D358')  # light yellow
    tcPr.append(shd)
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Inches(0.18)
    table.allow_autofit = True
    return table

def add_contact_box(doc, name, phone, email, website):
    doc.add_heading('Ihre Ansprechpartnerin', level=2)
    p = doc.add_paragraph()
    p.add_run(name + '\n').bold = True
    p.add_run(f'Tel: {phone}\n')
    p.add_run(f'E-Mail: {email}\n')
    p.add_run(f'Web: {website}')
    return p

def add_key_facts_table(doc, facts):
    table = doc.add_table(rows=1, cols=len(facts))
    row = table.rows[0]
    for i, (k, v) in enumerate(facts.items()):
        cell = row.cells[i]
        cell.text = f"{k}:\n{v}"
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.style = "Table Grid"
    return table

def add_vertical_facts_table(doc, facts_vertical):
    table = doc.add_table(rows=len(facts_vertical), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(facts_vertical):
        table.cell(i, 0).text = str(k)
        table.cell(i, 1).text = str(v)
    return table

def save_to_docx_with_images(text, logo_path=None, main_image_path=None,
                             detail_image_folder=None, output_dir="output",
                             title_prefix="anzeige", contact_info=None,
                             key_facts=None, specific_title=None, address=None,
                             call_to_action=None, contact_fields=None, facts_vertical=None):
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{title_prefix}-{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc = Document()
    # --- PAGE 1 ---
    doc.add_heading("Ihr Immobilienangebot", level=0)
    add_logo_top_right(doc, logo_path)
    add_main_image(doc, main_image_path, "")
    if specific_title:
        h = doc.add_heading(clean_heading(specific_title), level=1)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if address:
        p = doc.add_paragraph(address)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if key_facts:
        add_key_facts_table(doc, key_facts)
    if call_to_action:
        add_call_to_action(doc, call_to_action)
    if contact_fields:
        add_contact_box(doc, **contact_fields)
    doc.add_page_break()
    if detail_image_folder:
        add_image_gallery_from_folder(doc, detail_image_folder)
        doc.add_page_break()
    cleaned = clean_markdown(text)
    # --- DESCRIPTION (robust section detection) ---
    in_desc = False
    in_facts = False
    table_rows = []
    desc_found = False
    facts_found = False
    added_desc = False
    added_facts = False
    for elem_type, content in cleaned:
        heading = content.lower().strip() if isinstance(content, str) else ''
        # Start description section
        if (elem_type in ("heading2", "heading3") and "objektbeschreibung" in heading):
            doc.add_heading(clean_heading(content), level=2)
            in_desc = True
            in_facts = False
            desc_found = True
            continue
        # Start facts section
        if (elem_type in ("heading2", "heading3") and "daten und fakten" in heading):
            doc.add_heading(clean_heading(content), level=2)
            in_desc = False
            in_facts = True
            facts_found = True
            continue
        if in_desc:
            added_desc = True
            if elem_type == "paragraph":
                p = doc.add_paragraph()
                for t, val in content:
                    run = p.add_run(val)
                    run.bold = (t == 'bold')
            elif elem_type == "bullet":
                p = doc.add_paragraph(style='List Bullet')
                for t, val in content[1]:
                    run = p.add_run(val)
                    run.bold = (t == 'bold')
        elif in_facts:
            added_facts = True
            if elem_type == "table_row":
                table_rows.append(content)
            elif table_rows:
                add_clean_table_to_docx(doc, table_rows)
                table_rows = []
            elif elem_type == "paragraph":
                p = doc.add_paragraph()
                for t, val in content:
                    run = p.add_run(val)
                    run.bold = (t == 'bold')
    # Fallback: If no explicit description section, include all content after images up to facts
    if not desc_found and not added_desc:
        in_desc = True
        for elem_type, content in cleaned:
            heading = content.lower().strip() if isinstance(content, str) else ''
            if (elem_type in ("heading2", "heading3") and "daten und fakten" in heading):
                break
            if elem_type in ("heading2", "heading3") and "objektbeschreibung" in heading:
                continue
            if elem_type == "paragraph":
                p = doc.add_paragraph()
                for t, val in content:
                    run = p.add_run(val)
                    run.bold = (t == 'bold')
            elif elem_type == "bullet":
                p = doc.add_paragraph(style='List Bullet')
                for t, val in content[1]:
                    run = p.add_run(val)
                    run.bold = (t == 'bold')
    # --- Always add a dedicated 'Daten und Fakten' page with all key facts ---
    doc.add_page_break()
    doc.add_heading("Daten und Fakten", level=2)
    if facts_vertical:
        add_vertical_facts_table(doc, facts_vertical)
    elif key_facts:
        add_key_facts_table(doc, key_facts)
    # --- Also add any additional facts from LLM output ---
    if not facts_found and not added_facts:
        table_rows = []
        in_tables = False
        for elem_type, content in cleaned:
            heading = content.lower().strip() if isinstance(content, str) else ''
            if (elem_type in ("heading2", "heading3") and ("objektbeschreibung" in heading or "daten und fakten" in heading)):
                in_tables = True if "daten und fakten" in heading else False
                continue
            if in_tables:
                if elem_type == "table_row":
                    table_rows.append(content)
                elif table_rows:
                    add_clean_table_to_docx(doc, table_rows)
                    table_rows = []
                elif elem_type == "paragraph":
                    p = doc.add_paragraph()
                    for t, val in content:
                        run = p.add_run(val)
                        run.bold = (t == 'bold')
        if table_rows:
            add_clean_table_to_docx(doc, table_rows)
    # --- AGB ---
    doc.add_page_break()
    agb_text = load_agb_text()
    lines = [line.strip() for line in agb_text.split("\n") if line.strip()]
    if lines:
        doc.add_heading(clean_heading(lines[0]), level=2)
        para = None
        for line in lines[1:]:
            if re.match(r"^\d+\. ", line):
                doc.add_heading(clean_heading(line), level=3)
                para = None
            else:
                if para is None or para.text:
                    para = doc.add_paragraph()
                para.add_run(line)
    doc.save(filepath)
    return filepath